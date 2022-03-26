import os
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageOps
import shutil

def convert_image(main_images_folder, formato, new_width=2272):
    if not os.path.isdir(main_images_folder): 
        raise NotADirectoryError(f'{main_images_folder} não existe.')
    
    temporary_folder = os.path.join(main_images_folder, 'temporary_folder')
    
    try:
        os.mkdir(temporary_folder)
    except FileExistsError as e:
        print(f'{temporary_folder} já existe.')

    for root, dirs, files in os.walk(main_images_folder):
        for file in files: 
            file_full_path = os.path.join(root, file)
            file_name, extension = os.path.splitext(file)

            converted_tag = '_CONVERTED'            

            new_file = file_name + converted_tag + extension
            new_file_full_path = os.path.join(temporary_folder, new_file)

            # if converted_tag in file_full_path:
            #     os.remove(file_full_path)
            #
            # continue

            if os.path.isfile(new_file_full_path):
                print(f'Arquivo {new_file_full_path} já existe.')
                continue

            if converted_tag in file_full_path:
                print('Imagem já convertida.')
                continue
            
            
            img_pillow = Image.open(file_full_path)
            

            width, height = img_pillow.size
            new_height = round((new_width * height) / width)

            new_image = img_pillow.resize(
                (new_width, new_height),
                Image.LANCZOS
            )

            new_image.save(
                new_file_full_path,
                optimize=True,
                quality=80,
                #exif=img_pillow.info['exif'] quando quer aquelas informações de câmera e tal
            )

            print(f'{file_full_path} convertido com sucesso!')
            new_image.close()
            img_pillow.close()
    
    #Jogando no word
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run()    
    
    for root, dirs, files in os.walk(temporary_folder):
        for file in files:
            image = os.path.join(temporary_folder, file)
                        
            img_pillow = Image.open(image)
            
            new_image = ImageOps.expand(img_pillow, border=5, fill='white')
            
            new_image.save(
                    image,
                    optimize=True,
                    quality=80
                )   
                    
            new_image.close()
            img_pillow.close()                    
            
            if formato==0:
                run.add_picture(image, width=Inches(2.99))
            else:
                run.add_picture(image, width=Inches(2.00))                
    
    
    document.save(os.path.join(main_images_folder,'formatadas.docx'))
    
   
    shutil.rmtree(temporary_folder)





    